"""
Resume Kombat AI — Resume Parser Service
Extracts structured data from PDF, DOCX, and TXT files
"""

import io
import logging
import re
from typing import Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

# ── Optional heavy imports ─────────────────────────────────────────────────────
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("pdfplumber not installed — PDF parsing disabled")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed — DOCX parsing disabled")

try:
    import spacy
    try:
        NLP = spacy.load("en_core_web_sm")
        SPACY_AVAILABLE = True
    except OSError:
        SPACY_AVAILABLE = False
        logger.warning("spaCy en_core_web_sm not found — run: python -m spacy download en_core_web_sm")
except ImportError:
    SPACY_AVAILABLE = False
    logger.warning("spaCy not installed — NER disabled")


# ── Section Keywords ───────────────────────────────────────────────────────────
SECTION_HEADERS = {
    "experience":     ["experience", "work experience", "employment", "work history", "professional experience", "career"],
    "education":      ["education", "academic", "qualifications", "degrees", "university", "college"],
    "skills":         ["skills", "technical skills", "technologies", "tech stack", "competencies", "expertise", "proficiencies"],
    "projects":       ["projects", "personal projects", "portfolio", "side projects", "open source"],
    "certifications": ["certifications", "certificates", "credentials", "licenses", "accreditations"],
    "achievements":   ["achievements", "awards", "honors", "recognition", "accomplishments"],
    "summary":        ["summary", "objective", "profile", "about", "overview", "bio"],
}

# ── Common Tech Skills ─────────────────────────────────────────────────────────
TECH_SKILLS = {
    # Languages
    "python", "javascript", "typescript", "java", "go", "golang", "rust", "c++", "c#",
    "ruby", "php", "swift", "kotlin", "scala", "r", "matlab", "bash", "shell",
    # Frontend
    "react", "react.js", "reactjs", "next.js", "nextjs", "vue", "vue.js", "angular",
    "svelte", "html", "html5", "css", "css3", "tailwind", "bootstrap", "jquery",
    "webpack", "vite", "redux", "graphql",
    # Backend
    "node.js", "nodejs", "express", "fastapi", "django", "flask", "spring", "laravel",
    "rails", "gin", "fiber", "nestjs", "grpc",
    # Databases
    "postgresql", "mysql", "mongodb", "redis", "elasticsearch", "cassandra",
    "dynamodb", "sqlite", "neo4j", "influxdb", "bigquery", "snowflake",
    # Cloud
    "aws", "gcp", "azure", "cloudflare", "vercel", "netlify", "heroku",
    # DevOps
    "docker", "kubernetes", "k8s", "terraform", "ansible", "jenkins",
    "github actions", "gitlab ci", "circleci", "argocd", "helm", "istio",
    # ML/AI
    "tensorflow", "pytorch", "scikit-learn", "keras", "hugging face", "pandas",
    "numpy", "spark", "kafka", "airflow", "mlflow", "langchain",
    # Other
    "git", "linux", "microservices", "rest", "api", "ci/cd", "agile", "scrum",
}

SENIORITY_KEYWORDS = {
    "senior": 5, "staff": 7, "principal": 8, "lead": 6, "director": 9,
    "vp": 10, "head": 8, "chief": 10, "architect": 7, "manager": 6,
    "junior": 1, "intern": 0, "associate": 2, "mid": 3,
}


class ResumeParser:
    """Parse PDF/DOCX/TXT resumes into structured JSON"""

    def parse_bytes(self, content: bytes, filename: str) -> Dict:
        """Entry point: detect file type and extract text"""
        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else "txt"

        if ext == "pdf":
            text = self._parse_pdf(content)
        elif ext in ("docx", "doc"):
            text = self._parse_docx(content)
        else:
            text = content.decode("utf-8", errors="replace")

        return self.parse_text(text)

    def parse_text(self, text: str) -> Dict:
        """Parse raw text into structured resume data"""
        text = self._clean_text(text)
        sections = self._split_sections(text)

        result = {
            "name":                   self._extract_name(text),
            "email":                  self._extract_email(text),
            "phone":                  self._extract_phone(text),
            "linkedin":               self._extract_url(text, "linkedin"),
            "github":                 self._extract_url(text, "github"),
            "summary":                sections.get("summary", ""),
            "skills":                 self._extract_skills(text, sections.get("skills", "")),
            "experience":             self._extract_experience(sections.get("experience", "")),
            "education":              self._extract_education(sections.get("education", "")),
            "projects":               self._extract_projects(sections.get("projects", "")),
            "certifications":         self._extract_list(sections.get("certifications", "")),
            "achievements":           self._extract_list(sections.get("achievements", "")),
            "total_years_experience": 0.0,
            "seniority_level":        "mid",
        }

        result["total_years_experience"] = self._calc_years(result["experience"])
        result["seniority_level"] = self._calc_seniority(text, result["total_years_experience"])
        return result

    # ── Text Extraction ──────────────────────────────────────────────────────
    def _parse_pdf(self, content: bytes) -> str:
        if not PDF_AVAILABLE:
            return content.decode("utf-8", errors="replace")
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
                return "\n".join(pages)
        except Exception as e:
            logger.error(f"PDF parse error: {e}")
            return ""

    def _parse_docx(self, content: bytes) -> str:
        if not DOCX_AVAILABLE:
            return content.decode("utf-8", errors="replace")
        try:
            doc = DocxDocument(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            return "\n".join(paragraphs)
        except Exception as e:
            logger.error(f"DOCX parse error: {e}")
            return ""

    def _clean_text(self, text: str) -> str:
        # Normalize whitespace but preserve line structure
        text = re.sub(r"\r\n", "\n", text)
        text = re.sub(r"\r", "\n", text)
        text = re.sub(r"\t", " ", text)
        text = re.sub(r"[ ]{3,}", "  ", text)
        text = re.sub(r"\n{4,}", "\n\n\n", text)
        return text.strip()

    # ── Section Splitting ────────────────────────────────────────────────────
    def _split_sections(self, text: str) -> Dict[str, str]:
        sections = {}
        lines = text.split("\n")
        current_section = "header"
        current_lines = []

        for line in lines:
            stripped = line.strip().lower()
            matched = self._match_section(stripped)
            if matched and len(stripped) < 60:
                if current_lines:
                    sections[current_section] = "\n".join(current_lines)
                current_section = matched
                current_lines = []
            else:
                current_lines.append(line)

        if current_lines:
            sections[current_section] = "\n".join(current_lines)

        return sections

    def _match_section(self, line: str) -> Optional[str]:
        for section, keywords in SECTION_HEADERS.items():
            for kw in keywords:
                if kw in line and len(line) < len(kw) + 20:
                    return section
        return None

    # ── Field Extractors ─────────────────────────────────────────────────────
    def _extract_name(self, text: str) -> str:
        first_lines = text.strip().split("\n")[:5]
        for line in first_lines:
            line = line.strip()
            if 3 < len(line) < 60 and not re.search(r"[@|/\\:\.{}\[\]]", line):
                words = line.split()
                if 2 <= len(words) <= 5 and all(w[0].isupper() for w in words if w):
                    return line
        if SPACY_AVAILABLE:
            doc = NLP(text[:500])
            for ent in doc.ents:
                if ent.label_ == "PERSON":
                    return ent.text
        return ""

    def _extract_email(self, text: str) -> str:
        match = re.search(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", text)
        return match.group(0) if match else ""

    def _extract_phone(self, text: str) -> str:
        match = re.search(
            r"(\+?1?\s?)?(\(?\d{3}\)?[\s.\-]?\d{3}[\s.\-]?\d{4})", text
        )
        return match.group(0).strip() if match else ""

    def _extract_url(self, text: str, domain: str) -> str:
        pattern = rf"https?://(?:www\.)?{domain}\.com/[^\s]+"
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(0)
        # Also match bare handles
        handle_pattern = rf"{domain}\.com/([a-zA-Z0-9\-_/]+)"
        match = re.search(handle_pattern, text, re.IGNORECASE)
        return match.group(0) if match else ""

    def _extract_skills(self, full_text: str, skills_section: str) -> List[str]:
        source = (skills_section + "\n" + full_text).lower()
        found = set()
        for skill in TECH_SKILLS:
            # Word boundary matching
            pattern = rf"\b{re.escape(skill)}\b"
            if re.search(pattern, source):
                found.add(skill.title() if len(skill) > 3 else skill.upper())
        return sorted(found)

    def _extract_experience(self, text: str) -> List[Dict]:
        if not text.strip():
            return []
        items = []
        # Split by date patterns that signal new job entries
        job_blocks = re.split(
            r"\n(?=\S.{0,80}(?:20\d{2}|19\d{2}|present|current))",
            text, flags=re.IGNORECASE
        )
        for block in job_blocks[:10]:  # max 10 jobs
            if len(block.strip()) < 20:
                continue
            lines = [l.strip() for l in block.strip().split("\n") if l.strip()]
            if not lines:
                continue
            title = lines[0]
            company = lines[1] if len(lines) > 1 else ""
            bullets = [l for l in lines[2:] if l.startswith(("•", "-", "·", "*", "–")) or len(l) > 40]
            dates = re.findall(r"(20\d{2}|19\d{2}|present|current)", block, re.IGNORECASE)
            tech = [s.title() for s in TECH_SKILLS if s in block.lower()]
            items.append({
                "company": company,
                "title": title,
                "start_date": dates[0] if dates else "",
                "end_date": dates[1] if len(dates) > 1 else "Present",
                "duration_months": self._estimate_duration(dates),
                "bullets": bullets[:8],
                "technologies": tech[:10],
            })
        return items

    def _extract_education(self, text: str) -> List[Dict]:
        if not text.strip():
            return []
        items = []
        degree_keywords = ["b.s", "b.a", "m.s", "m.a", "phd", "ph.d", "mba",
                           "bachelor", "master", "doctorate", "associate", "diploma"]
        blocks = re.split(r"\n{2,}", text.strip())
        for block in blocks[:5]:
            lower = block.lower()
            if any(kw in lower for kw in degree_keywords + ["university", "college", "institute", "school"]):
                lines = [l.strip() for l in block.split("\n") if l.strip()]
                gpa_match = re.search(r"gpa[:\s]+(\d\.\d)", block, re.IGNORECASE)
                year_match = re.search(r"(20\d{2}|19\d{2})", block)
                items.append({
                    "institution": lines[0] if lines else "",
                    "degree": lines[1] if len(lines) > 1 else "",
                    "field": "",
                    "graduation_year": int(year_match.group(1)) if year_match else None,
                    "gpa": float(gpa_match.group(1)) if gpa_match else None,
                })
        return items

    def _extract_projects(self, text: str) -> List[Dict]:
        if not text.strip():
            return []
        items = []
        blocks = re.split(r"\n{2,}", text.strip())
        for block in blocks[:8]:
            lines = [l.strip() for l in block.split("\n") if l.strip()]
            if not lines:
                continue
            gh_match = re.search(r"github\.com/([^\s]+)", block, re.IGNORECASE)
            url_match = re.search(r"https?://[^\s]+", block)
            stars_match = re.search(r"(\d+[\.,]?\d*)[k\s]*stars?", block, re.IGNORECASE)
            tech = [s.title() for s in TECH_SKILLS if s in block.lower()]
            items.append({
                "name": lines[0],
                "description": lines[1] if len(lines) > 1 else "",
                "technologies": tech[:8],
                "github_url": gh_match.group(0) if gh_match else "",
                "url": url_match.group(0) if url_match else "",
                "stars": self._parse_stars(stars_match.group(1)) if stars_match else 0,
            })
        return items

    def _extract_list(self, text: str) -> List[str]:
        if not text.strip():
            return []
        items = []
        for line in text.split("\n"):
            line = line.strip().lstrip("•-·*–").strip()
            if 3 < len(line) < 120:
                items.append(line)
        return items[:15]

    # ── Calculations ─────────────────────────────────────────────────────────
    def _calc_years(self, experience: List[Dict]) -> float:
        total_months = sum(e.get("duration_months", 12) for e in experience)
        return round(total_months / 12, 1)

    def _estimate_duration(self, dates: List[str]) -> int:
        if len(dates) >= 2:
            try:
                start = int(dates[0])
                end = 2024 if dates[1].lower() in ("present", "current") else int(dates[1])
                return max(1, (end - start) * 12)
            except ValueError:
                pass
        return 12

    def _parse_stars(self, s: str) -> int:
        s = s.replace(",", "").replace(".", "")
        try:
            n = int(s)
            return n * 1000 if n < 100 else n
        except ValueError:
            return 0

    def _calc_seniority(self, text: str, years: float) -> str:
        lower = text.lower()
        for level, weight in sorted(SENIORITY_KEYWORDS.items(), key=lambda x: -x[1]):
            if level in lower:
                return level
        if years >= 8: return "senior"
        if years >= 4: return "mid"
        if years >= 1: return "junior"
        return "intern"

    def calc_ats_score(self, parsed: Dict, text: str) -> int:
        """Score resume ATS-friendliness 0-100"""
        score = 0
        # Contact info completeness (20 pts)
        if parsed.get("name"):      score += 5
        if parsed.get("email"):     score += 5
        if parsed.get("phone"):     score += 5
        if parsed.get("linkedin"):  score += 3
        if parsed.get("github"):    score += 2
        # Skills section (20 pts)
        skills = parsed.get("skills", [])
        score += min(20, len(skills) * 2)
        # Experience bullets (20 pts)
        exp = parsed.get("experience", [])
        bullet_count = sum(len(e.get("bullets", [])) for e in exp)
        score += min(20, bullet_count * 2)
        # Action verbs (10 pts)
        verbs = ["led", "built", "designed", "implemented", "improved", "reduced",
                 "increased", "launched", "managed", "created", "developed", "scaled"]
        verb_count = sum(1 for v in verbs if v in text.lower())
        score += min(10, verb_count)
        # Quantified achievements (15 pts)
        numbers = len(re.findall(r"\d+[%x]|\$\d+|\d+\s*(users|customers|ms|seconds|%)", text, re.IGNORECASE))
        score += min(15, numbers * 3)
        # Education present (5 pts)
        if parsed.get("education"):  score += 5
        # Certifications (5 pts)
        if parsed.get("certifications"):  score += 5
        # Reasonable length (5 pts)
        word_count = len(text.split())
        if 300 <= word_count <= 1000:  score += 5
        return min(100, score)


# Singleton
parser = ResumeParser()